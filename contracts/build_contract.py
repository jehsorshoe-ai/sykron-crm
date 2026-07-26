from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"E:\Codex\contracts\Modelo-Contrato-Prestacao-Servicos-Sykron.docx"
LOGO = r"E:\Codex\social\instagram\profile-sykron.png"

NAVY = "061321"
CYAN = "00B9F2"
BLUE = "2E74B5"
TEXT = "172A3A"
MUTED = "667788"
LIGHT = "F2F4F7"
WARN = "FFF4D6"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = tcMar.find(qn(f'w:{m}'))
        if el is None: el = OxmlElement(f'w:{m}'); tcMar.append(el)
        el.set(qn('w:w'), str(v)); el.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i])); tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)

def set_font(run, name='Calibri', size=11, bold=None, color=TEXT, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def add_p(doc, text='', bold_lead=None, after=6, keep=False, indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10; p.paragraph_format.left_indent = Inches(indent)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    if keep: p.paragraph_format.keep_with_next = True
    return p

def add_clause(doc, number, title):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"CLÁUSULA {number} - {title.upper()}"); set_font(r, size=13, bold=True, color=BLUE)
    return p

def add_section_heading(doc, label, title, page_break_before=False):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    r = p.add_run(f"{label} - {title.upper()}"); set_font(r, size=13, bold=True, color=BLUE)
    return p

def add_sub(doc, text):
    p = doc.add_paragraph(style='Heading 2'); p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_font(r, size=11.5, bold=True, color=TEXT)
    return p

def add_field_table(doc, rows):
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'; set_table_widths(t, [2300, 7060])
    for label, value in rows:
        cells = t.add_row().cells
        set_cell_shading(cells[0], LIGHT)
        p = cells[0].paragraphs[0]; r=p.add_run(label); set_font(r, bold=True, color=NAVY)
        p = cells[1].paragraphs[0]; r=p.add_run(value); set_font(r)
        for c in cells: c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(TEXT)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for name, size, color, before, after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,'1F4D78',8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header = sec.header
ht = header.add_table(rows=1, cols=2, width=Inches(6.5)); set_table_widths(ht,[1600,7760])
ht.cell(0,0).paragraphs[0].add_run().add_picture(LOGO, width=Inches(.42))
hp=ht.cell(0,1).paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=hp.add_run('SYKRON | MODELO CONTRATUAL'); set_font(r,size=9,bold=True,color=MUTED)

footer = sec.footer
fp = footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=fp.add_run('Modelo editável - preencher campos entre [colchetes] e revisar antes da assinatura | '); set_font(r,size=8,color=MUTED)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Masthead
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('SYKRON'); set_font(r,size=13,bold=True,color=CYAN)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
r=p.add_run('CONTRATO DE PRESTAÇÃO DE SERVIÇOS'); set_font(r,size=22,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(18)
r=p.add_run('Tecnologia, processos, sistemas, automação, inteligência artificial e indicadores'); set_font(r,size=10.5,italic=True,color=MUTED)

add_p(doc, 'Pelo presente instrumento particular, as partes abaixo identificadas:', align=WD_ALIGN_PARAGRAPH.LEFT)
add_field_table(doc, [
    ('CONTRATADA', '[RAZÃO SOCIAL / NOME EMPRESARIAL DA SYKRON]'),
    ('CNPJ/CPF', '[●]'), ('Endereço', '[●]'),
    ('Representante', '[NOME, CPF E CARGO]'), ('E-mail', 'contato.sykron@gmail.com'),
    ('CONTRATANTE', '[RAZÃO SOCIAL / NOME COMPLETO]'), ('CNPJ/CPF', '[●]'),
    ('Endereço', '[●]'), ('Representante', '[NOME, CPF E CARGO]'), ('E-mail', '[●]')
])
add_p(doc, 'têm entre si justo e contratado o presente Contrato de Prestação de Serviços, regido pelas cláusulas seguintes, pelo Anexo I (Escopo/Proposta Comercial) e, quando aplicável, pelo Anexo II (Tratamento de Dados Pessoais).')

add_clause(doc,'1','Objeto e documentos integrantes')
add_p(doc,'1.1. A CONTRATADA prestará os serviços profissionais descritos no Anexo I, que poderá abranger diagnóstico e melhoria de processos, implantação ou integração de sistemas, desenvolvimento de automações, painéis e indicadores, soluções com inteligência artificial, consultoria, treinamento e suporte.')
add_p(doc,'1.2. Integram este Contrato, em ordem de prevalência: (a) aditivos assinados; (b) este instrumento; (c) Anexo II; (d) Anexo I; e (e) propostas, ordens de serviço ou comunicações formalmente aprovadas por ambas as partes.')
add_p(doc,'1.3. Qualquer atividade não prevista no escopo será considerada serviço adicional e dependerá de aprovação escrita quanto a prazo, preço e impactos.')

add_clause(doc,'2','Escopo, entregáveis e aceite')
add_p(doc,'2.1. O Anexo I deverá identificar, no mínimo, objetivos, entregáveis, premissas, exclusões, cronograma, responsáveis, critérios de aceite e preço.')
add_p(doc,'2.2. A CONTRATANTE terá [5] dias úteis após a entrega para aprovar ou rejeitar fundamentadamente o entregável. A ausência de manifestação após lembrete escrito caracterizará aceite provisório, sem afastar vícios ocultos demonstráveis.')
add_p(doc,'2.3. Havendo rejeição fundamentada por desconformidade com o Anexo I, a CONTRATADA corrigirá o item sem custo adicional dentro de prazo razoável acordado. Mudanças de preferência, ampliação de escopo ou fatos externos serão tratados como solicitação de mudança.')

add_clause(doc,'3','Obrigações da contratada')
add_p(doc,'3.1. Executar os serviços com diligência técnica, boa-fé, pessoal qualificado e observância do escopo aprovado.')
add_p(doc,'3.2. Informar riscos, impedimentos e dependências relevantes; manter registros razoáveis das entregas; e proteger credenciais e informações recebidas.')
add_p(doc,'3.3. Não assumir obrigações, contratar terceiros em nome da CONTRATANTE ou acessar ambientes não autorizados.')
add_p(doc,'3.4. A CONTRATADA poderá utilizar subcontratados sob sua responsabilidade, preservadas confidencialidade, segurança e proteção de dados, salvo vedação expressa no Anexo I.')

add_clause(doc,'4','Obrigações da contratante')
add_p(doc,'4.1. Fornecer informações, acessos, decisões, homologações e profissionais necessários nos prazos acordados, garantindo que possui legitimidade para disponibilizar dados, sistemas e conteúdos.')
add_p(doc,'4.2. Designar responsável com poder de decisão; validar entregáveis; manter backups; e pagar os valores nas condições contratadas.')
add_p(doc,'4.3. Atrasos decorrentes de omissão, informação incorreta, indisponibilidade de ambiente ou demora de aprovação da CONTRATANTE prorrogarão o cronograma pelo período impactado, sem responsabilidade da CONTRATADA.')

add_clause(doc,'5','Preço, faturamento e tributos')
add_p(doc,'5.1. Pela execução dos serviços, a CONTRATANTE pagará: [VALOR, PERIODICIDADE, MARCOS E FORMA DE PAGAMENTO], conforme Anexo I.')
add_p(doc,'5.2. Salvo indicação diversa, despesas de viagem, licenças, APIs, nuvem, ferramentas e serviços de terceiros não estão incluídas e dependerão de autorização prévia.')
add_p(doc,'5.3. O atraso sujeitará o valor a multa moratória de 2%, juros de 1% ao mês pro rata die e correção monetária pelo IPCA, sem prejuízo da suspensão dos serviços após aviso escrito de [5] dias úteis.')
add_p(doc,'5.4. Cada parte será responsável pelos tributos que a legislação lhe atribuir. Retenções legais deverão ser comprovadas.')

add_clause(doc,'6','Prazo, cronograma e vigência')
add_p(doc,'6.1. O Contrato vigorará por [PRAZO] a partir de [DATA], podendo ser prorrogado por escrito. O cronograma do Anexo I considera o fornecimento tempestivo das dependências da CONTRATANTE.')
add_p(doc,'6.2. Serviços recorrentes poderão ser renovados por iguais períodos, desde que o Anexo I assim preveja, admitida denúncia sem justa causa com antecedência mínima de [30] dias.')

add_clause(doc,'7','Propriedade intelectual e licenças')
add_p(doc,'7.1. Permanecem de titularidade de cada parte os materiais, marcas, códigos, métodos, modelos, bibliotecas, ferramentas, conhecimentos e ativos preexistentes ao Contrato.')
add_p(doc,'7.2. Após o pagamento integral, a CONTRATANTE receberá os direitos sobre os entregáveis especificamente produzidos para ela na extensão definida no Anexo I. Na falta de definição, receberá licença não exclusiva, perpétua, irrevogável e mundial para uso interno dos entregáveis finais.')
add_p(doc,'7.3. Componentes genéricos, reutilizáveis, conhecimentos, técnicas, prompts, modelos, conectores, frameworks e melhorias de uso geral permanecem da CONTRATADA, assegurada à CONTRATANTE licença necessária ao uso do entregável.')
add_p(doc,'7.4. Software livre e componentes de terceiros observarão suas próprias licenças. Nenhuma parte poderá remover avisos de autoria ou utilizar marcas da outra sem autorização escrita.')

add_clause(doc,'8','Confidencialidade')
add_p(doc,'8.1. Informação Confidencial é toda informação técnica, comercial, estratégica, financeira, operacional ou pessoal identificada como confidencial ou que, pelas circunstâncias, deva razoavelmente ser protegida.')
add_p(doc,'8.2. A parte receptora usará a Informação Confidencial apenas para executar este Contrato, limitará o acesso a pessoas que necessitem conhecê-la e adotará medidas de proteção compatíveis com sua sensibilidade.')
add_p(doc,'8.3. Não são confidenciais informações que a receptora comprove: já conhecer legitimamente; serem públicas sem violação; terem sido recebidas licitamente de terceiro; ou terem sido desenvolvidas de forma independente.')
add_p(doc,'8.4. Se houver obrigação legal de divulgação, a parte receptora, quando permitido, notificará previamente a outra e limitará a divulgação ao estritamente exigido. A obrigação vigorará durante o Contrato e por [5] anos, sem limite para segredos de negócio enquanto mantiverem essa natureza.')

add_clause(doc,'9','Proteção de dados pessoais e segurança')
add_p(doc,'9.1. As partes cumprirão a Lei nº 13.709/2018 (LGPD). Quando a CONTRATADA tratar dados pessoais em nome da CONTRATANTE, esta será Controladora e aquela Operadora, salvo situação em que cada parte determine autonomamente as finalidades e meios de tratamento.')
add_p(doc,'9.2. Como Operadora, a CONTRATADA tratará dados somente conforme instruções documentadas e finalidades do Anexo II, restringirá acessos, manterá medidas técnicas e administrativas razoáveis e auxiliará a CONTRATANTE, dentro do escopo, no atendimento a titulares e autoridades.')
add_p(doc,'9.3. Incidente de segurança com risco ou dano relevante será comunicado à outra parte sem demora injustificada, com informações disponíveis sobre natureza, dados afetados, medidas de contenção e plano de resposta. A comunicação não implica reconhecimento de culpa.')
add_p(doc,'9.4. Suboperadores e transferências internacionais deverão observar a LGPD e garantias aplicáveis. Encerrado o tratamento, os dados serão devolvidos ou eliminados, ressalvadas obrigações legais de retenção e backups com ciclo regular de descarte.')

add_clause(doc,'10','Inteligência artificial e decisões automatizadas')
add_p(doc,'10.1. Quando houver uso de inteligência artificial, o Anexo I indicará finalidade, fontes, restrições, revisão humana e ferramentas de terceiros. Saídas de IA podem conter inexatidões e deverão ser validadas antes de uso relevante.')
add_p(doc,'10.2. A CONTRATANTE não inserirá dados pessoais sensíveis, segredos ou conteúdo protegido em ferramentas não aprovadas. Nenhuma solução será usada como única base para decisão jurídica, médica, financeira, trabalhista ou de alto impacto sem validação humana qualificada.')

add_clause(doc,'11','Garantias, suporte e níveis de serviço')
add_p(doc,'11.1. A CONTRATADA garante que os serviços serão executados de acordo com o Anexo I. Garantias específicas, período de correção, suporte, disponibilidade e tempos de resposta constarão do Anexo I.')
add_p(doc,'11.2. Salvo promessa expressa, a CONTRATADA não garante resultado econômico específico, ausência absoluta de falhas, compatibilidade com alterações futuras de terceiros ou disponibilidade contínua de serviços externos.')

add_clause(doc,'12','Responsabilidade')
add_p(doc,'12.1. Cada parte responderá pelos danos diretos que causar por descumprimento contratual comprovado, observado o nexo causal e o dever da parte prejudicada de mitigar o próprio prejuízo.')
add_p(doc,'12.2. Salvo dolo, fraude, violação de confidencialidade, propriedade intelectual, proteção de dados ou obrigação de pagamento, a responsabilidade total de cada parte ficará limitada ao valor efetivamente pago ou devido nos [12] meses anteriores ao evento que originou a reclamação.')
add_p(doc,'12.3. Na máxima extensão permitida por lei, nenhuma parte responderá por lucros cessantes, perda de oportunidade, danos indiretos ou perda de dados quando evitável por backup adequado, exceto se decorrentes de dolo ou quando a limitação for legalmente inaplicável.')

add_clause(doc,'13','Rescisão e efeitos')
add_p(doc,'13.1. O Contrato poderá ser rescindido: (a) por acordo; (b) por denúncia conforme Cláusula 6; (c) por inadimplemento não sanado em [10] dias úteis após notificação; ou (d) imediatamente em caso de ilicitude, violação grave de confidencialidade, insolvência ou risco relevante à segurança.')
add_p(doc,'13.2. Na rescisão, serão devidos os serviços executados, despesas autorizadas e compromissos não canceláveis. A CONTRATADA entregará os trabalhos pagos no estado em que se encontrarem e cooperará em transição adicional mediante contratação.')
add_p(doc,'13.3. Sobrevivem as cláusulas de pagamento, propriedade intelectual, confidencialidade, proteção de dados, responsabilidade, solução de controvérsias e demais que, por natureza, devam permanecer.')

add_clause(doc,'14','Ausência de vínculo e não exclusividade')
add_p(doc,'14.1. Este Contrato não cria vínculo empregatício, sociedade, representação, mandato, franquia ou exclusividade. Cada parte mantém autonomia técnica, administrativa e econômica e responde por seus profissionais e encargos.')

add_clause(doc,'15','Comunicações, alterações e assinatura eletrônica')
add_p(doc,'15.1. Notificações serão enviadas aos e-mails indicados na qualificação, considerando-se recebidas no primeiro dia útil após confirmação eletrônica ou ausência de mensagem de falha.')
add_p(doc,'15.2. Alterações de escopo, preço, prazo ou responsabilidade exigem documento escrito aceito por representantes autorizados. Aprovações operacionais poderão ocorrer por e-mail ou plataforma indicada no Anexo I.')
add_p(doc,'15.3. As partes reconhecem como válidas assinaturas eletrônicas que permitam comprovar autoria e integridade, inclusive por plataforma de assinatura aceita por ambas, nos termos da legislação aplicável, podendo este instrumento ser assinado em vias eletrônicas independentes.')

add_clause(doc,'16','Disposições gerais')
add_p(doc,'16.1. A tolerância não implica renúncia. A nulidade de uma disposição não prejudicará as demais, devendo as partes substituí-la por disposição válida de efeito econômico equivalente.')
add_p(doc,'16.2. Nenhuma parte poderá ceder integralmente o Contrato sem anuência da outra, exceto em reorganização societária que preserve capacidade de cumprimento. Caso fortuito ou força maior suspenderá obrigações afetadas enquanto perdurar o impedimento, com comunicação e mitigação.')

add_clause(doc,'17','Lei aplicável e foro')
add_p(doc,'17.1. Aplica-se a legislação brasileira. As partes buscarão solução amigável por [15] dias antes de medida judicial, salvo urgência. Fica eleito o foro da Comarca de [CIDADE/UF], com renúncia a qualquer outro, observado eventual foro legal obrigatório.')

add_p(doc,'E, por estarem de acordo, as partes assinam este instrumento, juntamente com duas testemunhas, físicas ou eletrônicas.', after=12)
add_p(doc,'[CIDADE/UF], [DIA] de [MÊS] de [ANO].', align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

sig=doc.add_table(rows=2,cols=2); set_table_widths(sig,[4680,4680])
for i,(title,detail) in enumerate([
    ('CONTRATADA','[RAZÃO SOCIAL]\nPor: [●]\nCargo: [●]'),
    ('CONTRATANTE','[RAZÃO SOCIAL]\nPor: [●]\nCargo: [●]'),
    ('TESTEMUNHA 1','Nome: [●]\nCPF: [●]'),
    ('TESTEMUNHA 2','Nome: [●]\nCPF: [●]')]):
    c=sig.cell(i//2,i%2); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('\n__________________________________\n'+title+'\n'); set_font(r,bold=True,color=NAVY)
    r=p.add_run(detail); set_font(r,size=9.5,color=MUTED)
doc.add_page_break()

add_section_heading(doc,'ANEXO I','Escopo / Ordem de Serviço')
add_field_table(doc,[
    ('Projeto','[NOME DO PROJETO]'),('Objetivo','[RESULTADO DE NEGÓCIO PRETENDIDO]'),
    ('Início','[DATA]'),('Prazo','[●]'),('Valor','[●]'),('Pagamento','[●]'),
    ('Gestor CONTRATANTE','[NOME/E-MAIL]'),('Gestor CONTRATADA','[NOME/E-MAIL]')
])
for title,body in [
('1. Escopo incluído','[Descrever atividades, sistemas, processos, integrações, automações, dashboards, treinamentos e suporte.]'),
('2. Entregáveis e critérios de aceite','[Listar cada entregável, formato, data, responsável por aprovação e critério verificável.]'),
('3. Exclusões','[Listar expressamente itens fora do preço e do prazo.]'),
('4. Premissas e dependências','[Acessos, qualidade de dados, equipe do cliente, licenças, ambientes, aprovações e terceiros.]'),
('5. Cronograma e marcos de pagamento','[Inserir marcos, datas e percentuais/valores.]'),
('6. Suporte e SLA, se aplicável','[Canais, horário, severidades, resposta, solução, disponibilidade e período de garantia.]'),
('7. Propriedade intelectual específica','[Definir cessão ou licença, código-fonte, documentação, componentes preexistentes e terceiros.]'),
('8. Solicitações de mudança','[Responsáveis, formato de aprovação e regra de reorçamento/replanejamento.]')]:
    add_sub(doc,title); add_p(doc,body,after=9)
add_section_heading(doc,'ANEXO II','Tratamento de Dados Pessoais (quando aplicável)', page_break_before=True)
add_field_table(doc,[
('Controladora','CONTRATANTE, salvo indicação diversa'),('Operadora','CONTRATADA, quando atuar sob instruções'),
('Finalidade','[●]'),('Titulares','[clientes, usuários, empregados, fornecedores etc.]'),
('Dados pessoais','[●]'),('Dados sensíveis','[não / especificar]'),('Duração','[●]'),
('Suboperadores','[●]'),('Transferência internacional','[não / países e garantias]'),
('Retenção e descarte','[●]'),('Contato de privacidade','[●]')
])
add_p(doc,'A CONTRATANTE declara possuir base legal e transparência adequadas para o tratamento e fornecerá instruções lícitas. A CONTRATADA manterá registro compatível com seu papel, controles de acesso, confidencialidade, gestão de incidentes e cooperação razoável para direitos de titulares. Requisitos adicionais de segurança, auditoria ou certificação deverão ser descritos neste Anexo e precificados no Anexo I.', after=0)

add_section_heading(doc,'NOTAS','Orientações de uso e referências legais')
t=doc.add_table(rows=1,cols=2); t.style='Table Grid'; set_table_widths(t,[2500,6860])
for i,h in enumerate(['Tema','Referência/uso']): set_cell_shading(t.rows[0].cells[i],NAVY); r=t.rows[0].cells[i].paragraphs[0].add_run(h); set_font(r,bold=True,color='FFFFFF')
refs=[
('Prestação de serviços','Código Civil, arts. 593 a 609. Ajustar o contrato ao serviço efetivamente prestado.'),
('Executividade','CPC, art. 784, III: documento particular assinado pelo devedor e por duas testemunhas pode constituir título executivo extrajudicial, se a obrigação for líquida, certa e exigível.'),
('Assinatura eletrônica','MP 2.200-2/2001, art. 10, §2º: admite outros meios de comprovação de autoria e integridade quando aceitos pelas partes.'),
('Dados pessoais','LGPD: definir papéis, finalidade, instruções, segurança, suboperadores, incidentes, retenção e direitos de titulares.'),
('Uso seguro','Preencher todos os campos [●], anexar escopo verificável, confirmar poderes dos signatários e revisar limites de responsabilidade, tributos, foro e regras setoriais com advogado brasileiro.')]
for a,b in refs:
    row=t.add_row().cells
    for cell in row:
        set_cell_margins(cell, top=25, start=75, bottom=25, end=75)
    r=row[0].paragraphs[0].add_run(a); set_font(r,size=8.4,bold=True,color=NAVY)
    r=row[1].paragraphs[0].add_run(b); set_font(r,size=8)

doc.core_properties.title='Modelo de Contrato de Prestação de Serviços - Sykron'
doc.core_properties.subject='Modelo B2B brasileiro para serviços de tecnologia, processos, automação e IA'
doc.core_properties.author='Sykron'
doc.save(OUT)
print(OUT)
